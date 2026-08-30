#!/usr/bin/env python3
"""
DocMerge 3.0 for Mac
Combines .txt, .doc, .docx, .pdf, .rtf and .md files into a single Word
(.docx) document. Runs entirely on device; no files leave your machine.

UI deliberately avoids ttk.Treeview and ttk styling, which are broken on the
Tk 8.5 that ships with Apple's system Python. Everything here is plain Tk
(Listbox, Label, Canvas) which renders correctly on every macOS Python.

Built for Steve Derviniotis.
"""

from __future__ import annotations

import os
import queue
import subprocess
import sys
import tempfile
import threading
import traceback
from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_BREAK
from pypdf import PdfReader

SUPPORTED = {".txt", ".doc", ".docx", ".pdf", ".rtf", ".md"}
APP_NAME = "DocMerge"
VERSION = "3.1"

TYPE_LABELS = {
    ".txt": "TXT", ".md": "MD", ".doc": "DOC",
    ".docx": "DOCX", ".pdf": "PDF", ".rtf": "RTF",
}


def fmt_size(b: float) -> str:
    for unit in ("B", "KB", "MB", "GB"):
        if b < 1024:
            return f"{b:.0f} {unit}" if unit == "B" else f"{b:.1f} {unit}"
        b /= 1024
    return f"{b:.1f} TB"


# ============================ Engine ============================ #

def _read_text(path: Path) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return path.read_text(encoding=enc)
        except (UnicodeDecodeError, UnicodeError):
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def extract_txt(path: Path):
    return _read_text(path).split("\n")


def extract_doc_via_textutil(path: Path):
    with tempfile.TemporaryDirectory() as tmp:
        out = Path(tmp) / (path.stem + ".txt")
        result = subprocess.run(
            ["textutil", "-convert", "txt", "-output", str(out), str(path)],
            capture_output=True, text=True)
        if result.returncode != 0 or not out.exists():
            raise RuntimeError(
                f"textutil could not convert this file: "
                f"{result.stderr.strip() or 'unknown error'}")
        return extract_txt(out)


def extract_docx(path: Path):
    doc = Document(str(path))
    lines = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            lines.append(" | ".join(c.text.strip() for c in row.cells))
    return lines


def extract_pdf(path: Path, page_cb=None):
    reader = PdfReader(str(path))
    total = len(reader.pages)
    lines = []
    for i, page in enumerate(reader.pages, start=1):
        if page_cb:
            page_cb(i, total)
        text = page.extract_text() or ""
        if text.strip():
            lines.extend(text.split("\n"))
        else:
            lines.append(
                f"[Page {i}: no extractable text; likely a scanned image]")
        if i < total:
            lines.append("")
    return lines


def extract_file(path: Path, page_cb=None):
    ext = path.suffix.lower()
    if ext in (".txt", ".md"):
        return extract_txt(path)
    if ext == ".docx":
        return extract_docx(path)
    if ext in (".doc", ".rtf"):
        try:
            return extract_doc_via_textutil(path)
        except (FileNotFoundError, RuntimeError):
            try:
                return extract_docx(path)
            except Exception:
                raise RuntimeError(
                    "Could not read this legacy .doc; textutil conversion "
                    "failed.")
    if ext == ".pdf":
        return extract_pdf(path, page_cb=page_cb)
    raise ValueError(f"Unsupported file type: {ext}")


@dataclass
class MergeReport:
    merged: int = 0
    skipped: list = field(default_factory=list)
    output: object = None
    out_size: int = 0


def build_merged_docx(files, output: Path, status_cb=None, progress_cb=None):
    def status(msg):
        if status_cb:
            status_cb(msg)

    def progress(frac):
        if progress_cb:
            progress_cb(max(0.0, min(1.0, frac)))

    n = len(files)
    status("Preparing document")
    progress(0.02)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    doc.add_heading("Merged Document", level=0)
    meta = doc.add_paragraph()
    run = meta.add_run(f"{n} source files")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    report = MergeReport()
    progress(0.05)

    for idx, path in enumerate(files):
        base = 0.05 + 0.85 * (idx / n)
        span = 0.85 / n
        label = f"({idx + 1} of {n})"
        status(f"Reading {path.name} {label}")
        progress(base)

        def page_cb(p, total, _b=base, _s=span, _nm=path.name, _l=label):
            status(f"Extracting {_nm} {_l} — page {p} of {total}")
            progress(_b + _s * 0.7 * (p / total))

        try:
            lines = extract_file(path, page_cb=page_cb)
        except Exception as e:
            report.skipped.append(f"{path.name}: {e}")
            continue

        status(f"Adding {path.name} to document {label}")
        progress(base + span * 0.8)

        brk = doc.add_paragraph()
        brk.add_run().add_break(WD_BREAK.PAGE)
        doc.add_heading(path.name, level=1)
        for line in lines:
            doc.add_paragraph(line)

        report.merged += 1
        progress(base + span)

    status("Saving Word document")
    progress(0.93)
    doc.save(str(output))
    report.output = output
    try:
        report.out_size = output.stat().st_size
    except OSError:
        report.out_size = 0
    status("Done")
    progress(1.0)
    return report


# ============================ GUI ============================ #

class DocMergeGUI:
    # Dark palette. Buttons use ttk with the "clam" theme, because plain
    # tk.Button ignores background colour on macOS Aqua (which made v2's
    # buttons render as white text on a native light button).
    BG = "#1e2430"
    CARD = "#252d3b"
    ROW_ALT = "#2a3342"
    LINE = "#3a4457"
    INK = "#e8ecf2"
    SUBTLE = "#98a3b5"
    ACCENT = "#4f8cff"
    ACCENT_DK = "#3a6fd8"
    BTN = "#333d4f"
    BTN_HOVER = "#3f4a5f"
    OK = "#4fc08d"
    WARN = "#e6b455"

    def __init__(self, root):
        self.root = root
        self.files = []
        self.msg_q = queue.Queue()
        self.merging = False
        self.last_output = None

        root.title(APP_NAME)
        root.geometry("720x640")
        root.minsize(660, 560)
        root.configure(bg=self.BG)

        self._init_styles()
        self._build()
        self._poll()
        self._refresh()

    # ---------- widget helpers ----------

    def _init_styles(self):
        """ttk with the 'clam' theme is the only reliable way to get coloured
        buttons on macOS; the native Aqua theme discards background colour."""
        st = ttk.Style()
        try:
            st.theme_use("clam")
        except tk.TclError:
            pass

        st.configure("Tool.TButton", background=self.BTN, foreground=self.INK,
                     bordercolor=self.LINE, focuscolor=self.BTN,
                     lightcolor=self.BTN, darkcolor=self.BTN,
                     relief="flat", padding=(9, 6),
                     font=("Helvetica", 11))
        st.map("Tool.TButton",
               background=[("pressed", self.ACCENT_DK),
                           ("active", self.BTN_HOVER),
                           ("disabled", "#2a3342")],
               foreground=[("disabled", self.SUBTLE),
                           ("pressed", "#ffffff")])

        st.configure("Primary.TButton", background=self.ACCENT,
                     foreground="#ffffff", bordercolor=self.ACCENT,
                     focuscolor=self.ACCENT, lightcolor=self.ACCENT,
                     darkcolor=self.ACCENT, relief="flat",
                     padding=(20, 10), font=("Helvetica", 13, "bold"))
        st.map("Primary.TButton",
               background=[("pressed", self.ACCENT_DK),
                           ("active", self.ACCENT_DK),
                           ("disabled", "#3a4358")],
               foreground=[("disabled", self.SUBTLE)])

        st.configure("Vertical.TScrollbar", background=self.BTN,
                     troughcolor=self.CARD, bordercolor=self.CARD,
                     arrowcolor=self.SUBTLE)

    def _button(self, parent, text, cmd, primary=False, width=None):
        b = ttk.Button(parent, text=text, command=cmd,
                       style="Primary.TButton" if primary else "Tool.TButton")
        if width:
            b.config(width=width)
        return b

    def _build(self):
        # ---- Header ----
        head = tk.Frame(self.root, bg=self.BG)
        head.pack(fill="x", padx=18, pady=(14, 2))
        tk.Label(head, text=APP_NAME, bg=self.BG, fg=self.INK,
                 font=("Helvetica", 20, "bold")).pack(side="left")
        tk.Label(head, text=f"v{VERSION}", bg=self.BG, fg=self.SUBTLE,
                 font=("Helvetica", 11)).pack(side="left", padx=(8, 0),
                                              pady=(8, 0))
        tk.Label(head, text="Runs on this Mac. Nothing is uploaded.",
                 bg=self.BG, fg=self.SUBTLE,
                 font=("Helvetica", 10)).pack(side="right", pady=(8, 0))

        # ---- Toolbar ----
        bar = tk.Frame(self.root, bg=self.BG)
        bar.pack(fill="x", padx=18, pady=(8, 6))
        for text, cmd, pad in (
            ("Add Files…", self.add_files, (0, 6)),
            ("Add Folder…", self.add_folder, (0, 14)),
            ("Remove", self.remove_selected, (0, 6)),
            ("↑ Up", lambda: self.move(-1), (0, 6)),
            ("↓ Down", lambda: self.move(1), (0, 6)),
        ):
            self._button(bar, text, cmd).pack(side="left", padx=pad)
        self._button(bar, "Clear", self.clear_all).pack(side="right")

        # ---- Column header ----
        colhead = tk.Frame(self.root, bg=self.BG)
        colhead.pack(fill="x", padx=18)
        tk.Label(colhead, text="  #   FILE NAME", bg=self.BG, fg=self.SUBTLE,
                 font=("Helvetica", 9, "bold"), anchor="w").pack(side="left")
        tk.Label(colhead, text="TYPE        SIZE  ", bg=self.BG,
                 fg=self.SUBTLE, font=("Helvetica", 9, "bold"),
                 anchor="e").pack(side="right")

        # ---- File list (plain Listbox: reliable on every Tk) ----
        wrap = tk.Frame(self.root, bg=self.LINE, bd=1, relief="solid")
        wrap.pack(fill="both", expand=True, padx=18, pady=(2, 0))
        self.listbox = tk.Listbox(
            wrap, selectmode="extended", activestyle="none",
            bg=self.CARD, fg=self.INK,
            selectbackground=self.ACCENT, selectforeground="#ffffff",
            highlightthickness=0, bd=0, font=("Menlo", 12))
        sb = ttk.Scrollbar(wrap, orient="vertical",
                           command=self.listbox.yview)
        self.listbox.config(yscrollcommand=sb.set)
        self.listbox.pack(side="left", fill="both", expand=True)
        sb.pack(side="right", fill="y")
        self.listbox.bind("<<ListboxSelect>>", lambda e: self._refresh_sel())

        # ---- Summary strip ----
        summ = tk.Frame(self.root, bg=self.BG)
        summ.pack(fill="x", padx=18, pady=(8, 0))
        self.summary_lbl = tk.Label(summ, text="", bg=self.BG, fg=self.INK,
                                    font=("Helvetica", 12, "bold"),
                                    anchor="w")
        self.summary_lbl.pack(side="left")
        self.sel_lbl = tk.Label(summ, text="", bg=self.BG, fg=self.SUBTLE,
                                font=("Helvetica", 11), anchor="e")
        self.sel_lbl.pack(side="right")

        # ---- Status line ----
        self.stage_lbl = tk.Label(self.root, text="Ready", bg=self.BG,
                                  fg=self.SUBTLE, anchor="w",
                                  font=("Helvetica", 11))
        self.stage_lbl.pack(fill="x", padx=18, pady=(6, 2))

        # ---- Canvas progress bar (hand-drawn: no ttk) ----
        self.pb_h = 14
        self.pbar = tk.Canvas(self.root, height=self.pb_h, bg=self.BG,
                              highlightthickness=0, bd=0)
        self.pbar.pack(fill="x", padx=18)
        self.pbar.bind("<Configure>", lambda e: self._draw_progress())
        self._progress = 0.0

        # ---- Result panel ----
        self.result_frame = tk.Frame(self.root, bg=self.BG)
        self.result_lbl = tk.Label(self.result_frame, text="", bg=self.BG,
                                   fg=self.OK, anchor="w", justify="left",
                                   font=("Helvetica", 11, "bold"))
        self.result_lbl.pack(side="left")
        self.reveal_btn = self._button(self.result_frame, "Show in Finder",
                                       self.reveal)
        self.reveal_btn.pack(side="right")

        # ---- Footer ----
        foot = tk.Frame(self.root, bg=self.BG)
        self.foot = foot
        foot.pack(side="bottom", fill="x", padx=18, pady=(10, 14))
        tk.Label(foot,
                 text=f"Python {sys.version_info.major}."
                      f"{sys.version_info.minor}  ·  Tk {tk.TkVersion}",
                 bg=self.BG, fg=self.SUBTLE,
                 font=("Helvetica", 9)).pack(side="left")
        self.merge_btn = self._button(foot, "Merge to Word Document",
                                      self.start_merge, primary=True)
        self.merge_btn.pack(side="right", ipadx=10, ipady=2)

    # ---------- progress drawing ----------

    def _draw_progress(self):
        c = self.pbar
        c.delete("all")
        w = c.winfo_width() or 600
        h = self.pb_h
        c.create_rectangle(0, 3, w, h - 3, fill="#333d4f", outline="")
        if self._progress > 0:
            fill_w = max(3, int(w * self._progress))
            colour = self.OK if self._progress >= 1.0 else self.ACCENT
            c.create_rectangle(0, 3, fill_w, h - 3, fill=colour, outline="")

    def set_progress(self, frac):
        self._progress = frac
        self._draw_progress()

    # ---------- list rendering ----------

    def _row_text(self, i, p):
        try:
            size = fmt_size(p.stat().st_size)
        except OSError:
            size = "?"
        kind = TYPE_LABELS.get(p.suffix.lower(), "?")
        name = p.name
        if len(name) > 42:
            name = name[:39] + "…"
        return f" {i + 1:>2}  {name:<42} {kind:>5} {size:>10}"

    def _refresh(self):
        self.listbox.delete(0, "end")
        if not self.files:
            self.listbox.insert("end", "")
            self.listbox.insert(
                "end", "     No files added yet.")
            self.listbox.insert(
                "end", "     Click  Add Files…  or  Add Folder…  above.")
            self.listbox.itemconfig(1, foreground=self.SUBTLE)
            self.listbox.itemconfig(2, foreground=self.SUBTLE)
            self.summary_lbl.config(text="0 files selected")
            self.sel_lbl.config(text="")
            return

        total = 0
        for i, p in enumerate(self.files):
            self.listbox.insert("end", self._row_text(i, p))
            if i % 2 == 1:
                self.listbox.itemconfig(i, background=self.ROW_ALT)
            try:
                total += p.stat().st_size
            except OSError:
                pass

        kinds = {}
        for p in self.files:
            k = TYPE_LABELS.get(p.suffix.lower(), "?")
            kinds[k] = kinds.get(k, 0) + 1
        breakdown = ", ".join(f"{v} {k}" for k, v in sorted(kinds.items()))
        n = len(self.files)
        self.summary_lbl.config(
            text=f"{n} file{'s' if n != 1 else ''} ready  ·  "
                 f"{fmt_size(total)} total  ·  {breakdown}")
        self._refresh_sel()

    def _refresh_sel(self):
        sel = len(self.listbox.curselection()) if self.files else 0
        self.sel_lbl.config(text=f"{sel} selected" if sel else "")

    # ---------- file management ----------

    def add_files(self):
        paths = filedialog.askopenfilenames(
            title="Select files to merge",
            filetypes=[("Supported files",
                        "*.txt *.doc *.docx *.pdf *.rtf *.md"),
                       ("All files", "*.*")])
        added, rejected = 0, 0
        for p in paths:
            p = Path(p)
            if p.suffix.lower() not in SUPPORTED:
                rejected += 1
            elif p in self.files:
                rejected += 1
            else:
                self.files.append(p)
                added += 1
        self._refresh()
        msg = f"Added {added} file{'s' if added != 1 else ''}"
        if rejected:
            msg += f"; {rejected} skipped (duplicate or unsupported type)"
        self.stage_lbl.config(text=msg, fg=self.INK)
        self._hide_result()

    def add_folder(self):
        folder = filedialog.askdirectory(title="Select a folder")
        if not folder:
            return
        found = sorted(p for p in Path(folder).iterdir()
                       if p.suffix.lower() in SUPPORTED
                       and not p.name.startswith(("~$", ".")))
        new = [p for p in found if p not in self.files]
        if not new:
            messagebox.showinfo(
                APP_NAME, "No new supported files found in that folder.")
            return
        self.files.extend(new)
        self._refresh()
        self.stage_lbl.config(
            text=f"Added {len(new)} files from “{Path(folder).name}”",
            fg=self.INK)
        self._hide_result()

    def remove_selected(self):
        if not self.files:
            return
        idxs = sorted(self.listbox.curselection(), reverse=True)
        if not idxs:
            self.stage_lbl.config(text="Select a row first, then Remove",
                                  fg=self.WARN)
            return
        for i in idxs:
            if 0 <= i < len(self.files):
                del self.files[i]
        self._refresh()
        self.stage_lbl.config(text=f"Removed {len(idxs)} file(s)",
                              fg=self.INK)

    def clear_all(self):
        self.files.clear()
        self._refresh()
        self.set_progress(0)
        self._hide_result()
        self.stage_lbl.config(text="Ready", fg=self.SUBTLE)

    def move(self, direction):
        if not self.files:
            return
        sel = sorted(self.listbox.curselection())
        if not sel:
            self.stage_lbl.config(text="Select a row first, then move it",
                                  fg=self.WARN)
            return
        if direction < 0 and sel[0] == 0:
            return
        if direction > 0 and sel[-1] >= len(self.files) - 1:
            return
        order = sel if direction < 0 else list(reversed(sel))
        new_sel = []
        for i in order:
            j = i + direction
            self.files[i], self.files[j] = self.files[j], self.files[i]
            new_sel.append(j)
        self._refresh()
        for j in new_sel:
            self.listbox.selection_set(j)
        self._refresh_sel()

    # ---------- result panel ----------

    def _show_result(self, text, colour):
        self.result_lbl.config(text=text, fg=colour)
        self.result_frame.pack(side="bottom", fill="x", padx=18,
                               pady=(10, 0), after=self.foot)

    def _hide_result(self):
        self.result_frame.pack_forget()

    def reveal(self):
        if self.last_output and sys.platform == "darwin":
            subprocess.run(["open", "-R", str(self.last_output)])

    # ---------- merge ----------

    def start_merge(self):
        if self.merging:
            return
        if not self.files:
            messagebox.showinfo(APP_NAME, "Add some files first.")
            return
        output = filedialog.asksaveasfilename(
            title="Save merged document as",
            defaultextension=".docx",
            initialfile="Merged Document.docx",
            filetypes=[("Word Document", "*.docx")])
        if not output:
            return

        self.merging = True
        self.merge_btn.config(state="disabled", text="Merging…")
        self._hide_result()
        self.set_progress(0)
        files = list(self.files)

        def worker():
            try:
                r = build_merged_docx(
                    files, Path(output),
                    status_cb=lambda m: self.msg_q.put(("status", m)),
                    progress_cb=lambda f: self.msg_q.put(("progress", f)))
                self.msg_q.put(("done", r))
            except Exception:
                self.msg_q.put(("error", traceback.format_exc()))

        threading.Thread(target=worker, daemon=True).start()

    def _poll(self):
        try:
            while True:
                kind, payload = self.msg_q.get_nowait()
                if kind == "status":
                    self.stage_lbl.config(text=payload, fg=self.INK)
                elif kind == "progress":
                    self.set_progress(payload)
                elif kind == "done":
                    self._done(payload)
                elif kind == "error":
                    self._error(payload)
        except queue.Empty:
            pass
        self.root.after(60, self._poll)

    def _done(self, r):
        self.merging = False
        self.merge_btn.config(state="normal", text="Merge to Word Document")
        self.set_progress(1.0)
        self.last_output = r.output

        name = Path(r.output).name
        if r.skipped:
            self.stage_lbl.config(
                text=f"Finished with {len(r.skipped)} skipped file(s)",
                fg=self.WARN)
            self._show_result(
                f"Merged {r.merged} of {r.merged + len(r.skipped)} files → "
                f"{name}  ({fmt_size(r.out_size)})", self.WARN)
            messagebox.showwarning(
                APP_NAME,
                f"Merged {r.merged} files into:\n{r.output}\n\n"
                f"Size: {fmt_size(r.out_size)}\n\nSkipped:\n"
                + "\n".join(f"  • {s}" for s in r.skipped))
        else:
            self.stage_lbl.config(text="Merge complete", fg=self.OK)
            self._show_result(
                f"Merged {r.merged} files → {name}  "
                f"({fmt_size(r.out_size)})", self.OK)
            messagebox.showinfo(
                APP_NAME,
                f"Merged {r.merged} files into:\n{r.output}\n\n"
                f"Size: {fmt_size(r.out_size)}")

        if sys.platform == "darwin" and r.output:
            subprocess.run(["open", "-R", str(r.output)])

    def _error(self, tb):
        self.merging = False
        self.merge_btn.config(state="normal", text="Merge to Word Document")
        self.stage_lbl.config(text="Merge failed", fg=self.WARN)
        self.set_progress(0)
        messagebox.showerror(APP_NAME, f"Merge failed:\n\n{tb}")


# ============================ Entry ============================ #

def main():
    if len(sys.argv) > 2:
        output = Path(sys.argv[1])
        files = [Path(p) for p in sys.argv[2:]]
        r = build_merged_docx(files, output,
                              status_cb=lambda m: print(f"  {m}"))
        print(f"Merged {r.merged} of {len(files)} files into {output} "
              f"({fmt_size(r.out_size)})")
        for s in r.skipped:
            print(f"  Skipped {s}")
        return

    global tk, ttk, filedialog, messagebox
    import tkinter as tk
    from tkinter import ttk, filedialog, messagebox

    root = tk.Tk()
    DocMergeGUI(root)
    root.mainloop()


if __name__ == "__main__":
    main()
